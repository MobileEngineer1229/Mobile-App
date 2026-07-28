"""Common module for extracting training text from files of various formats.

【Beginner's Guide】
  This file contains the training pipeline "data entrance" Everything.
  Read files of various formats and create strings(string)Change it to.
  preprocess.pyWow train_tokenizer.py uses this module.

  iter_texts() Process all formats with one function:
    for text in iter_texts(Path("data/raw")):
        # textis the string of each document

Supported Format
---------
    .txt            regular characters(UTF-8)
    .json           JSON Array or single object
    .jsonl          JSON row format (one per line JSON object)
    .pdf            PDF document (text extraction)
    .docx           Microsoft Word Document
    image file     .jpg .jpeg .png .bmp .gif .webp .tiff .tif (OCR character recognition)

Install dependent libraries (Just install the required format)
--------------------
    pip install pdfplumber     ← PDF support
    pip install python-docx    ← DOCX support
    pip install easyocr        ← image OCR support (about 200MB Automatic download of model)

How to use
---------
    from src.data.readers import iter_texts
    for text in iter_texts(Path("data/raw")):
        process(text)
"""

from __future__ import annotations

import json
import sys
from pathlib import Path
from typing import Iterator

# List of image supported extensions
IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png", ".bmp", ".gif", ".webp", ".tiff", ".tif"}

# EasyOCR Create an instance only once and reuse it (Reduce initialization costs)
_ocr_reader = None


def _get_ocr_reader():
    """EasyOCR returns an instance. The model is downloaded upon first call.."""
    global _ocr_reader
    if _ocr_reader is None:
        try:
            import easyocr
        except ImportError:
            raise ImportError(
                "Image file processing easyocris required.\n"
                "installation command: pip install easyocr"
            )
        print(
            "[readers] EasyOCR Initializing character recognizer...\n"
            "          Korean language when first run/Download the English recognition model (about 200MB).",
            file=sys.stderr,
        )
        # gpu=True: Enable graphics processing unit acceleration (If not, automatically cpu use)
        _ocr_reader = easyocr.Reader(["ko", "en"], gpu=True)
    return _ocr_reader


# ---------------------------------------------------------------------------
# Format-specific read functions
# ---------------------------------------------------------------------------

def _read_txt(path: Path) -> Iterator[str]:
    """UTF-8 Read plain text files."""
    try:
        content = path.read_text(encoding="utf-8", errors="replace").strip()
        if content:
            yield content
    except Exception as e:
        print(f"[readers] txt read error {path}: {e}", file=sys.stderr)


def _read_jsonl(path: Path) -> Iterator[str]:
    """JSON row format(.jsonl) reads the file.

    Supported Format:
        {"text": "..."}
        {"question": "...", "answer": "..."}
        {"content": "..."}
        {"body": "..."}
    """
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            for lineno, line in enumerate(f, 1):
                line = line.strip()
                if not line:
                    continue
                try:
                    obj = json.loads(line)
                except json.JSONDecodeError:
                    continue
                text = _extract_text_from_obj(obj)
                if text:
                    yield text
    except Exception as e:
        print(f"[readers] jsonl read error {path}: {e}", file=sys.stderr)


def _read_json(path: Path) -> Iterator[str]:
    """JSON file(.json)reads.

    support structure:
        single object:  {"text": "..."}
        arrangement:       [{"text": "..."}, ...]
        string array: ["sentence1", "sentence2", ...]
    """
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            data = json.load(f)
    except Exception as e:
        print(f"[readers] json read error {path}: {e}", file=sys.stderr)
        return

    if isinstance(data, list):
        for item in data:
            if isinstance(item, str) and item.strip():
                yield item.strip()
            elif isinstance(item, dict):
                text = _extract_text_from_obj(item)
                if text:
                    yield text
    elif isinstance(data, dict):
        text = _extract_text_from_obj(data)
        if text:
            yield text


def _read_pdf(path: Path) -> Iterator[str]:
    """PDF Extract text from a file.

    pdfplumber I use a library.
    installation: pip install pdfplumber
    """
    try:
        import pdfplumber
    except ImportError:
        print(
            f"[readers] PDF to processing pdfplumberis required.\n"
            f"          installation command: pip install pdfplumber\n"
            f"          Skip: {path}",
            file=sys.stderr,
        )
        return

    try:
        with pdfplumber.open(str(path)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        yield text.strip()
                except Exception as e:
                    print(f"[readers] PDF {path} {page_num}page error: {e}", file=sys.stderr)
    except Exception as e:
        print(f"[readers] PDF read error {path}: {e}", file=sys.stderr)


def _read_docx(path: Path) -> Iterator[str]:
    """Microsoft Word Document(.docx)Extract text from.

    python-docx I use a library.
    installation: pip install python-docx
    """
    try:
        from docx import Document
    except ImportError:
        print(
            f"[readers] DOCX to processing python-docxis required.\n"
            f"          installation command: pip install python-docx\n"
            f"          Skip: {path}",
            file=sys.stderr,
        )
        return

    try:
        doc = Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if paragraphs:
            yield "\n".join(paragraphs)
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        yield text
    except Exception as e:
        print(f"[readers] DOCX read error {path}: {e}", file=sys.stderr)


def _read_image(path: Path) -> Iterator[str]:
    """From image file OCRRecognize text with.

    EasyOCR I use a library (Joseon language + English support).
    installation: pip install easyocr

    caution: When you run it for the first time, the recognition model is automatically downloaded. (about 200MB).
    """
    try:
        reader = _get_ocr_reader()
    except ImportError as e:
        print(f"[readers] {e}\n          Skip: {path}", file=sys.stderr)
        return

    try:
        # detail=0: Return only letters (no coordinates), paragraph=True: Join lines
        results = reader.readtext(str(path), detail=0, paragraph=True)
        text = "\n".join(r for r in results if isinstance(r, str) and r.strip())
        if text:
            yield text
    except Exception as e:
        print(f"[readers] image OCR error {path}: {e}", file=sys.stderr)


# ---------------------------------------------------------------------------
# public JSON Object text extraction
# ---------------------------------------------------------------------------

def _extract_text_from_obj(obj: dict) -> str:
    """JSON Extract text from an object.

    Keys that recognize: text, question+answer, content, body, passage, context
    """
    if not isinstance(obj, dict):
        return ""

    # plain text field
    for key in ("text", "content", "body", "passage", "context", "sentence"):
        if key in obj and isinstance(obj[key], str) and obj[key].strip():
            return obj[key].strip()

    # question and answer format
    if "question" in obj and "answer" in obj:
        q = str(obj["question"]).strip()
        a = str(obj["answer"]).strip()
        if q and a:
            return f"question: {q}\nanswer: {a}"

    # title + body format
    if "title" in obj and "content" in obj:
        title = str(obj.get("title", "")).strip()
        content = str(obj.get("content", "")).strip()
        if content:
            return f"{title}\n{content}" if title else content

    return ""


# ---------------------------------------------------------------------------
# main entry point
# ---------------------------------------------------------------------------

def iter_texts(raw_dir: Path) -> Iterator[str]:
    """raw_dir Extract text from files in all supported formats below.

    Files are traversed recursively down to subfolders.
    Deduplication is done by the caller(preprocess.py)Handled by.
    """
    for path in sorted(raw_dir.rglob("*")):
        if not path.is_file():
            continue
        suffix = path.suffix.lower()

        if suffix == ".txt":
            yield from _read_txt(path)
        elif suffix == ".jsonl":
            yield from _read_jsonl(path)
        elif suffix == ".json":
            yield from _read_json(path)
        elif suffix == ".pdf":
            yield from _read_pdf(path)
        elif suffix == ".docx":
            yield from _read_docx(path)
        elif suffix in IMAGE_SUFFIXES:
            yield from _read_image(path)
        # Other formats(yes: .doc, .xlsx etc.)is quietly skipped
